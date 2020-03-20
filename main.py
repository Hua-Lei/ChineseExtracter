import re
import docx
from tkinter import *
import tkinter.messagebox
from tkinter.filedialog import askopenfilename

def Path():
    path1 = askopenfilename()
    path.set(path1)
def doc():
    filename=e2.get()
    e1name = e1.get()
    if e1name == "":
        tkinter.messagebox.showerror('错误','未填写源文件地址')
        return;
    if filename == "":
        tkinter.messagebox.showerror('错误','未填写输出文件名')
        return;
    doc = docx.Document(e1name)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    text = "".join(fulltext)
    textList = text.split()
    for i in textList:
        res1 = "".join(re.findall("[\u4e00-\u9fa5]",i))
        if(res1 != ""):
            newdoc.add_paragraph(res1)
    newdoc.save("./result/" + filename)
    tkinter.messagebox.showinfo('提示','已完成提取操作')


newdoc = docx.Document()

root = Tk()
root.title('ChineseExtracter')
root.geometry('300x200')

path = StringVar()
name = StringVar()
l1 = Label(root,text = "抽取文件地址：")
l1.pack()
e1 = Entry(root,textvariable=path)
e1.pack()
btn1 = Button(root,text="选择文件路径",command = Path)
btn1.pack()
l2 = Label(root,text = "保存文件名（例如text.docx）")
l2.pack()
e2 = Entry(root)
e2.pack()
btn2 = Button(root,text="确认",command = doc)
btn2.pack()
l3 = Label(root,text = "         by 华泪HuaLei")
l3.pack()

root.mainloop()